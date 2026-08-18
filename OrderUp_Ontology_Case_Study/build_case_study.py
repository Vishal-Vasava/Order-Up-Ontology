from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "OrderUp_Ontology_Based_Platform_Case_Study.docx"
ARCH = ROOT / "orderup_future_architecture.png"

NAVY = "17324D"
BLUE = "245B78"
TEAL = "087F78"
ORANGE = "B45309"
GOLD = "D69E2E"
INK = "1F2937"
MUTED = "667085"
LIGHT = "F3F6F8"
PALE_BLUE = "EAF2F7"
PALE_TEAL = "E7F5F3"
PALE_ORANGE = "FFF3E6"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margin(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            cell_margin(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=None, bold=None, color=INK, italic=None, font="Aptos"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        a = p.add_run(bold_lead)
        set_run(a, bold=True)
        b = p.add_run(text[len(bold_lead):])
        set_run(b)
    else:
        r = p.add_run(text)
        set_run(r)
    return p


def add_bullets(doc, items, style="List Bullet"):
    for item in items:
        p = doc.add_paragraph(style=style)
        r = p.add_run(item)
        set_run(r)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run(r)


def add_callout(doc, title, text, fill=PALE_TEAL, color=TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title.upper())
    set_run(r, size=9, bold=True, color=color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_run(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths, header_fill=NAVY, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for i, text in enumerate(headers):
        shade(hdr.cells[i], header_fill)
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        set_run(r, size=9, bold=True, color=WHITE)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if ridx % 2:
                shade(cells[i], "F8FAFB")
            p = cells[i].paragraphs[0]
            r = p.add_run(str(value))
            set_run(r, size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def page_break(doc):
    doc.add_page_break()


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.82)
sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(1.0)
sec.right_margin = Inches(1.0)
sec.header_distance = Inches(0.4)
sec.footer_distance = Inches(0.4)

# Style system: standard_business_brief, with a named OrderUp palette override.
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in (
    ("Heading 1", 16, NAVY, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, TEAL, 8, 4),
):
    st = styles[name]
    st.font.name = "Aptos Display"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True
for lname in ("List Bullet", "List Number"):
    st = styles[lname]
    st.font.name = "Aptos"
    st.font.size = Pt(10.5)
    st.paragraph_format.left_indent = Inches(0.5)
    st.paragraph_format.first_line_indent = Inches(-0.25)
    st.paragraph_format.space_after = Pt(5)
    st.paragraph_format.line_spacing = 1.10

if "Figure Caption" not in [s.name for s in styles]:
    cap = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
else:
    cap = styles["Figure Caption"]
cap.font.name = "Aptos"
cap.font.size = Pt(9)
cap.font.italic = True
cap.font.color.rgb = RGBColor.from_string(MUTED)
cap.paragraph_format.space_before = Pt(4)
cap.paragraph_format.space_after = Pt(8)
cap.paragraph_format.keep_with_next = False

# Header/footer.
header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = header.add_run("ORDERUP  |  ONTOLOGY-BASED PLATFORM CASE STUDY")
set_run(hr, size=8, bold=True, color=MUTED)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run("Architecture research working paper  |  1 August 2026")
set_run(fr, size=8, color=MUTED)

# Cover.
doc.add_paragraph().paragraph_format.space_after = Pt(52)
k = doc.add_paragraph()
k.alignment = WD_ALIGN_PARAGRAPH.CENTER
kr = k.add_run("ORDERUP CASE STUDY")
set_run(kr, size=10, bold=True, color=ORANGE)
k.paragraph_format.space_after = Pt(10)
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = t.add_run("From Deterministic Ecommerce\nto an Ontology-Based AI Platform")
set_run(tr, size=28, bold=True, color=NAVY, font="Aptos Display")
t.paragraph_format.space_after = Pt(12)
s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = s.add_run("Future-state architecture, nondeterministic capabilities, LLM and voice strategy, governance, and migration roadmap")
set_run(sr, size=13, color=BLUE)
s.paragraph_format.space_after = Pt(30)
add_callout(doc, "Architecture thesis", "Preserve the existing Node.js and MySQL transaction system as the deterministic authority. Add an event-driven operational ontology, model platform, and governed action layer around it so AI can understand context, recommend decisions, and automate only within explicit policy boundaries.", fill=PALE_BLUE, color=BLUE)
doc.add_paragraph().paragraph_format.space_after = Pt(46)
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("Prepared as an AI architecture research case study\nRepository reviewed read-only; no application code changed\n1 August 2026")
set_run(mr, size=10, color=MUTED)

page_break(doc)

heading(doc, "Executive summary", 1)
add_body(doc, "OrderUp is a multi-persona ecommerce and fulfillment platform serving customers, store managers, delivery agents, and platform managers through Flutter mobile/web experiences and a React administration portal. Its Node.js/Express API and MySQL database implement the operational system of record for identity, catalog, inventory, carts, orders, delivery, telemetry, notifications, returns, claims, and reviews.")
add_body(doc, "The recommended future state is not to replace this core with a graph database or an LLM. Instead, OrderUp should build an operational ontology as a semantic digital twin of commerce activity, synchronized through reliable domain events. Specialized prediction models and LLMs can reason over that context, while a governed action gateway sends validated commands back through the existing API.")
add_callout(doc, "Recommended direction", "Build an ontology-enabled intelligence plane around the deterministic core. Let AI read broadly, reason probabilistically, and propose narrowly; let deterministic services validate and commit every consequential business transaction.", fill=PALE_TEAL, color=TEAL)

heading(doc, "Expected business value", 2)
add_bullets(doc, [
    "More relevant product discovery and contextual customer assistance.",
    "Better delivery assignment, routing, ETA prediction, and exception handling.",
    "Demand forecasting, replenishment support, and inventory-risk visibility for stores.",
    "Cross-domain operational intelligence for platform managers without manually joining many tables.",
    "A safe foundation for conversational and voice agents across all four personas.",
    "Faster evolution of AI models because semantic contracts are separated from transactional schemas.",
])

heading(doc, "Principal recommendation", 2)
add_body(doc, "Start with an order-store-inventory-delivery-temperature ontology. This domain has the strongest relationships in the present platform and supports valuable, measurable, comparatively low-risk use cases: ETA prediction, delivery-agent ranking, fulfillment exception detection, cold-chain risk, and operational root-cause analysis.")

heading(doc, "Scope and evidence", 1)
add_body(doc, "This case study is based on a read-only inspection of the provided OrderUp repository. The assessment considered the Flutter application, React administration portal, Node.js API, reverse-engineered MySQL schema, route/controller/model organization, order-state handling, delivery assignment, notifications, and temperature/location records. No source files, schemas, or runtime configuration were changed.")

heading(doc, "Current platform at a glance", 1)
add_table(doc, ["Area", "Observed implementation", "Architectural implication"], [
    ("Experiences", "Flutter mobile/web and React admin portal", "Shared APIs already form a practical migration boundary."),
    ("Personas", "Customer, store manager/producer, delivery agent, platform manager", "Ontology must represent roles independently from person identity."),
    ("Application", "Node.js and Express controllers, routes, models", "Retain as the deterministic command and transaction layer."),
    ("Operational data", "MySQL with users, products, inventory, carts, orders, returns, telemetry, notifications", "Strong source for a semantic read model; it should remain authoritative."),
    ("Identity and messaging", "Firebase authentication and push notifications", "Preserve integration; expose identity and consent context semantically."),
    ("Payments and assets", "Razorpay and object/image storage", "Wrap through existing deterministic services and governed tools."),
    ("Fulfillment", "Item-level status, delivery agent, truck, addresses, history", "Natural starting point for operational graph and decision models."),
], [1800, 3380, 4180])

heading(doc, "Current-state strengths", 2)
add_bullets(doc, [
    "A recognizable separation among clients, API, and operational database.",
    "Rich commerce and fulfillment data, including item-level status and historical observations.",
    "Four clearly identified operational personas with distinct workflows.",
    "Existing notification, location, and temperature channels that can support proactive intelligence.",
    "A small client-side ontology abstraction that demonstrates interest in typed entities and relationships.",
])

heading(doc, "Current-state limitations", 2)
add_bullets(doc, [
    "Business meaning is distributed across numeric codes, table names, SQL joins, controller logic, and client assumptions.",
    "The Flutter ontology abstraction is client-owned and is not a governed platform semantic layer.",
    "There is no durable representation for inferred facts, confidence, provenance, evidence, or model versions.",
    "Cross-domain questions require application-specific joins instead of stable semantic queries.",
    "Probabilistic recommendations and deterministic execution are not separated by a formal action boundary.",
    "Delivery assignment can fall back to a random selection rather than a transparent context-aware ranking.",
])

page_break(doc)
heading(doc, "Architecture principles", 1)
add_numbered(doc, [
    "Preserve the system of record. MySQL and existing commerce services remain authoritative for orders, inventory, money, identity, and status transitions.",
    "Build the ontology as a semantic read model. The graph describes business meaning and relationships; it does not become an alternate transaction authority.",
    "Synchronize with events, not controller dual writes. Use a transactional outbox or change-data capture so ontology projections can be replayed and reconciled.",
    "Represent uncertainty explicitly. Predictions and inferred facts carry confidence, evidence, model version, observation time, and expiration.",
    "Separate recommendation from execution. Models produce structured recommendations or action proposals; policies and deterministic services decide whether to execute.",
    "Use the right model for each job. LLMs handle language and reasoning; forecasting, ranking, anomaly detection, and optimization use specialized models.",
    "Design for hybrid inference. A model gateway should support cloud APIs, private local models, small models, and on-device models without coupling clients to a vendor.",
    "Make human oversight proportional to impact. Low-risk assistance can be automated earlier; financial, identity, and policy-sensitive decisions require confirmation or approval.",
])

heading(doc, "Future-state architecture", 1)
add_body(doc, "The target platform introduces an intelligence plane between user experiences and the deterministic commerce core. Ontology-aware retrieval gives models grounded context. A model gateway selects cloud or local inference. A governed action gateway converts approved proposals into typed commands executed by the existing Node.js API.")
if ARCH.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(ARCH), width=Inches(6.45))
    p.paragraph_format.keep_with_next = True
    cap_p = doc.add_paragraph(style="Figure Caption")
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.add_run("Figure 1. Proposed layered architecture: multimodal experiences and nondeterministic intelligence around the unchanged deterministic commerce core.")

heading(doc, "Layer responsibilities", 2)
add_table(doc, ["Layer", "Responsibility"], [
    ("Experience and personas", "Flutter, admin portal, operational dashboards, conversational copilots, and voice interaction."),
    ("Multimodal interaction", "API/BFF, sessions, realtime delivery, speech processing, vision, and document intake."),
    ("AI orchestration and models", "Context assembly, model routing, tool selection, generation, structured outputs, and evaluation."),
    ("Semantic access", "Ontology-aware RAG, stable domain queries, policy checks, and governed action proposals."),
    ("Ontology and intelligence", "Canonical concepts, graph relationships, temporal facts, vectors, features, reasoning, and predictions."),
    ("Event integration", "Outbox/CDC, versioned domain events, projectors, entity resolution, and reconciliation."),
    ("Deterministic commerce core", "Authentication, carts, orders, inventory, prices, payments, returns, fulfillment states, and authoritative persistence."),
], [2350, 7010])

page_break(doc)
heading(doc, "Operational ontology design", 1)
add_body(doc, "The ontology should model business meaning rather than mechanically convert each relational table into a node type. Transactional tables remain optimized for writes; the ontology is optimized for context, traversal, inference, explainability, and stable domain semantics.")

heading(doc, "Core concepts", 2)
add_table(doc, ["Domain", "Canonical concepts"], [
    ("Identity and organization", "Person, Customer, Worker, DeliveryAgent, StoreManager, PlatformManager, Role, Permission, Store, Organization, ServiceArea"),
    ("Catalog and inventory", "Product, ProductVariant, Category, UnitOfMeasure, CatalogListing, Price, Promotion, InventoryPosition, StockMovement"),
    ("Commerce", "Cart, CartItem, Order, OrderLine, Payment, Charge, Refund, Address, Review"),
    ("Fulfillment", "Fulfillment, Delivery, DeliveryTask, Route, Stop, AgentAssignment, Vehicle, DeliveryWindow, ProofOfDelivery"),
    ("Returns and service", "ReturnRequest, ReplacementRequest, Claim, Reason, Evidence, Resolution, Exception"),
    ("Telemetry", "LocationObservation, TemperatureObservation, Device, Sensor, Threshold, Breach"),
    ("Intelligence", "Recommendation, Prediction, RiskAssessment, Explanation, Decision, ActionProposal, PolicyEvaluation, HumanApproval, ActionExecution, ModelVersion"),
], [2350, 7010])

heading(doc, "Representative relationships", 2)
add_bullets(doc, [
    "Customer PLACED Order; Order CONTAINS OrderLine; OrderLine REFERENCES Product.",
    "Store OFFERS CatalogListing; CatalogListing LISTS Product; Store HOLDS InventoryPosition.",
    "OrderLine FULFILLED_BY Store; Order REQUIRES Fulfillment; Fulfillment ASSIGNED_TO DeliveryAgent.",
    "Fulfillment USES Vehicle; Delivery HAS_OBSERVATION TemperatureObservation; Claim CONCERNS OrderLine.",
    "Recommendation PROPOSES ActionProposal; Decision BASED_ON Evidence; ActionExecution AUTHORIZED_BY PolicyEvaluation.",
])

heading(doc, "Fact contract", 2)
add_body(doc, "Every asserted, inferred, or predicted fact should carry enough context to be trusted and retired safely. Required metadata should include source system and record, event or observation time, ingestion time, tenant/store scope, assertion type, confidence, model or rule version, evidence references, and expiration where applicable.")
add_callout(doc, "Temporal reasoning example", "“Customer prefers morning delivery” should not be stored as an eternal fact. It may be an inference with 0.76 confidence, supported by six recent orders, generated by preference-model version 3, and expiring after 90 days.", fill=PALE_ORANGE, color=ORANGE)

heading(doc, "Semantic actions", 2)
add_body(doc, "An operational ontology must describe what may be done, not only what exists. Candidate actions include RecommendProducts, ProposeCart, ReserveInventory, PlaceOrder, RecommendDeliveryAssignment, AssignDeliveryAgent, ProposeRoute, UpdateFulfillmentStatus, ReportDeliveryException, RecommendClaimResolution, ApproveRefund, CreatePromotionDraft, PublishPromotion, SendNotification, and EscalateCase.")
add_body(doc, "Each action contract should define the authorized proposer and approver, input schema, preconditions, deterministic validations, impact limits, human-approval requirement, idempotency key, compensation behavior, and audit requirements.")

heading(doc, "LLM and model architecture", 1)
add_body(doc, "A unified model gateway should insulate the platform from model vendors and deployment modes. The orchestrator requests a capability and operating constraints; the gateway selects the appropriate model based on quality, latency, privacy, availability, and cost.")
add_table(doc, ["Model class", "Best-fit responsibilities", "Deployment options"], [
    ("Frontier/general LLM", "Complex reasoning, conversation, synthesis, tool planning, multilingual assistance", "Cloud API or capable self-hosted model"),
    ("Small language model", "Intent, classification, extraction, routing, concise summarization", "Local service, edge, or on-device"),
    ("Embedding/reranker", "Semantic retrieval, similarity, candidate reranking", "Managed API or local inference"),
    ("Speech models", "Voice activity, denoising, speech recognition, language detection, speech generation", "Cloud, local, or device hybrid"),
    ("Vision model", "Product intake, claim evidence, proof-of-delivery quality, OCR-assisted understanding", "Cloud or private GPU"),
    ("Predictive model", "ETA, demand, churn, return risk, anomaly detection", "Feature-based managed or self-hosted service"),
    ("Optimization solver", "Routing, assignment, capacity allocation, replenishment", "Deterministic or stochastic optimization service"),
], [1700, 4560, 3100], font_size=8.8)

heading(doc, "Cloud versus local inference", 2)
add_table(doc, ["Choice", "Advantages", "Trade-offs", "Recommended use"], [
    ("Cloud LLM API", "Fast adoption; strong reasoning; managed scaling", "Network dependency; usage cost; governance and vendor considerations", "Complex copilots and early experimentation"),
    ("Local LLM", "Privacy control; predictable environment; custom deployment", "GPU operations; capacity; model lifecycle; potentially lower quality", "Sensitive operational contexts and stable high-volume tasks"),
    ("Small/on-device", "Low latency; offline support; low marginal cost", "Limited context and reasoning", "Voice commands, intent, extraction, and offline drafts"),
    ("Hybrid gateway", "Best model per task; graceful fallback; progressive portability", "Requires strong evaluation and routing policy", "Preferred strategic architecture"),
], [1300, 2500, 2800, 2760], font_size=8.5)

heading(doc, "Ontology-aware RAG", 2)
add_body(doc, "Retrieval should combine graph traversal, semantic search, feature access, and deterministic API lookups. A customer assistant needs only the subgraph and documents relevant to that customer, location, products, stores, and active order. A platform manager may receive broader aggregated context. Authorization must be applied before retrieval, not only after the model responds.")

heading(doc, "Voice interaction architecture", 1)
add_body(doc, "Voice is especially valuable for delivery agents because it supports hands-free work. The pipeline should include optional wake-word detection, voice-activity detection, noise suppression, speech-to-text, language and intent detection, ontology-grounded reasoning, governed action confirmation, and text-to-speech.")

heading(doc, "Example delivery-agent voice flow", 2)
add_numbered(doc, [
    "The delivery agent says, “Customer is unavailable. Delay this delivery.”",
    "The speech pipeline removes noise, transcribes the utterance, and records language and recognition confidence.",
    "The copilot retrieves the agent's assigned delivery, current route, order, customer instructions, and applicable policy.",
    "The copilot asks for explicit confirmation: “Delay by 15 minutes and notify the customer?”",
    "After confirmation, it submits typed ReportDeliveryException and UpdateETA proposals.",
    "The action gateway validates identity, assignment, permissions, impact, and policy before invoking the Node.js API.",
    "The API commits the transaction to MySQL, emits a domain event, and returns the confirmed outcome for spoken feedback.",
])
add_callout(doc, "Voice safety rule", "Never interpret silence as approval. Separate dictation from execution, require confirmation for consequential actions, preserve the transcript and action trail, and support queued drafts when connectivity is poor.", fill=PALE_ORANGE, color=ORANGE)

page_break(doc)
heading(doc, "Nondeterministic capabilities by persona", 1)
add_body(doc, "Nondeterministic features may be probabilistic, generative, ranking-based, optimization-based, or agentic. They should complement—not absorb—the deterministic commerce rules.")

heading(doc, "Customer", 2)
add_table(doc, ["Capability", "AI contribution", "Deterministic boundary"], [
    ("Semantic product discovery", "Interpret intent and rank products across catalog, store, availability, and location context", "API returns authoritative price and availability"),
    ("Conversational shopping", "Build a proposed cart and explain choices or substitutes", "Cart mutation and checkout use existing commands"),
    ("Personalized recommendations", "Rank products using preferences, purchases, reviews, and feasibility", "Eligibility, price, consent, and inventory remain rules"),
    ("ETA and proactive assistance", "Predict arrival ranges and identify likely delays or cold-chain risk", "Committed status and remedies follow policy"),
    ("Natural-language return", "Extract item, reason, evidence, and requested remedy", "Eligibility, refund, and approval remain deterministic"),
], [2150, 4210, 3000], font_size=8.7)

heading(doc, "Store manager", 2)
add_bullets(doc, [
    "Demand forecasting by product, store, weekday, season, and delivery area.",
    "Replenishment recommendations with demand, lead-time, and safety-stock explanations.",
    "Inventory anomaly, shrinkage, depletion, and stale-stock detection.",
    "Product onboarding assistance for extraction, categorization, descriptions, units, and tags.",
    "Promotion drafts based on inventory, margin, demand, and customer cohorts.",
    "Fulfillment workload forecasting and exception prioritization.",
    "Review and claim clustering into actionable store or product issues.",
])

heading(doc, "Delivery agent", 2)
add_bullets(doc, [
    "Context-aware assignment ranking using distance, workload, capacity, vehicle, skills, service area, and reliability.",
    "Route and stop-sequence optimization across pickups and deliveries.",
    "Continuous ETA prediction and customer-notification recommendations.",
    "Voice briefings, proposed status updates, and exception capture.",
    "Proof-of-delivery quality analysis and inconsistent-location detection.",
    "Predictive temperature-risk alerts before a threshold breach occurs.",
])

heading(doc, "Platform manager", 2)
add_bullets(doc, [
    "An ontology-powered operations command center that explains emerging issues across stores, orders, deliveries, inventory, and claims.",
    "Natural-language analytics over governed semantic questions.",
    "Fraud and abuse risk analysis connecting users, devices, addresses, returns, stores, and transactions.",
    "Claim triage, evidence assembly, and resolution recommendations.",
    "Store health and operational-risk scoring.",
    "Policy-impact simulation before fee, delivery, or service-area changes.",
    "Data-quality stewardship and ontology-governance assistance.",
])

heading(doc, "What must remain deterministic", 1)
add_table(doc, ["Always deterministic", "AI may assist by"], [
    ("Authentication, authorization, role and tenant enforcement", "Classifying intent and selecting permitted tools"),
    ("Price, fee, tax, discount, and promotion calculation", "Drafting or recommending a promotion"),
    ("Inventory reservation, decrement, and release", "Forecasting demand and proposing replenishment"),
    ("Order creation and allowed status transitions", "Summarizing context and proposing the next action"),
    ("Payment, capture, refund, and reconciliation", "Collecting evidence and recommending a resolution"),
    ("Return and claim eligibility", "Classifying reasons and estimating risk"),
    ("Consent, PII access, retention, and account actions", "Detecting sensitive data and prompting safe handling"),
    ("Idempotency, concurrency, and audit persistence", "Generating an idempotent typed action proposal"),
], [4550, 4810], font_size=8.8)

heading(doc, "Governed action pattern", 2)
add_numbered(doc, [
    "A model emits a schema-valid recommendation or action proposal with confidence, explanation, evidence, and model version.",
    "The action gateway authenticates the actor and evaluates role, tenant, resource ownership, policy, and impact limits.",
    "The system requests human approval when the action class or impact requires it.",
    "The gateway invokes an allowlisted existing API command with an idempotency key; it never sends raw SQL.",
    "The deterministic core validates all preconditions and commits or rejects the transaction.",
    "The outcome is recorded with the proposal, policy evaluation, approval, command result, and resulting domain event.",
])

heading(doc, "Governance and safety", 1)
add_table(doc, ["Control", "Purpose"], [
    ("Identity and tenant isolation", "Prevent cross-store or cross-customer context leakage."),
    ("Retrieval authorization", "Apply policy before graph, vector, or feature retrieval."),
    ("PII controls", "Minimize, redact, tokenize, and audit sensitive information."),
    ("Model and prompt registry", "Version prompts, tools, schemas, models, and routing policies."),
    ("Evaluation and red teaming", "Test grounding, tool selection, bias, refusal, prompt injection, and unsafe action behavior."),
    ("Confidence and impact thresholds", "Make autonomy proportional to model certainty and business consequence."),
    ("Human approval", "Provide accountable review for financial, identity, policy, and irreversible actions."),
    ("Tracing and audit", "Reconstruct evidence, reasoning inputs, policy checks, commands, and outcomes."),
    ("Quality, drift, and cost monitoring", "Detect degradation, distribution shifts, excessive latency, and unsustainable spend."),
], [2600, 6760])

heading(doc, "Automation levels", 2)
add_table(doc, ["Level", "Behavior", "OrderUp example"], [
    ("L0 Observe", "No recommendation or action", "Detect a likely late delivery"),
    ("L1 Recommend", "Rank or explain; user decides", "Recommend delivery agents"),
    ("L2 Draft", "Prepare a structured action for approval", "Draft a claim resolution or promotion"),
    ("L3 Bounded automation", "Execute only under strict policy, confidence, and impact limits", "Send a low-risk delay notification"),
    ("L4 Prohibited autonomy", "Unbounded or direct high-impact execution", "Arbitrary pricing, refunds, account changes, or database writes"),
], [1100, 4000, 4260], font_size=8.8)

page_break(doc)
heading(doc, "Migration roadmap", 1)
add_table(doc, ["Phase", "Primary outcome", "Representative work"], [
    ("0. Language and governance", "A shared semantic contract", "Domain vocabulary, owners, status semantics, identity, tenant, privacy, retention, ontology and event versioning"),
    ("1. Semantic read model", "An event-synchronized operational graph", "Transactional outbox/CDC, backfill, versioned events, projectors, entity resolution, reconciliation"),
    ("2. Semantic query", "Stable domain APIs and read-only intelligence", "Ontology service, query APIs, graph security, ontology-aware RAG, platform and store copilots"),
    ("3. Models and recommendations", "Measurable nondeterministic capabilities", "ETA, demand, anomaly, ranking, route optimization, feature store, feedback capture, shadow mode"),
    ("4. Governed actions", "Safe recommendations-to-execution boundary", "Action schemas, policy engine, approval workflow, idempotency, decision audit"),
    ("5. Controlled agents", "Multi-step workflows within explicit authority", "Exception resolution, replenishment, delivery orchestration, and claim investigation agents"),
], [1150, 2700, 5510], font_size=8.5)

heading(doc, "Recommended first releases", 2)
add_numbered(doc, [
    "Read-only platform operations copilot over orders, stores, inventory, delivery, claims, and telemetry.",
    "Store-manager inventory-risk and demand insights.",
    "Customer semantic product discovery with grounded availability and price retrieval.",
    "Delivery ETA prediction and at-risk fulfillment alerts.",
    "Context-aware delivery-agent ranking in shadow mode, compared with current assignment outcomes.",
    "Voice-based delivery exception reporting with explicit confirmation and no autonomous status commit.",
])

heading(doc, "Pilot recommendation: intelligent delivery assignment", 1)
add_body(doc, "Delivery assignment is the strongest first closed-loop case because the current platform already tracks agents, order items, locations, stores, addresses, and status. The new capability can initially run in shadow mode without changing assignments.")
add_table(doc, ["Stage", "Description", "Success evidence"], [
    ("Eligibility", "Deterministic filtering of active, authorized, available agents with valid service area and vehicle constraints", "No ineligible agent appears in candidates"),
    ("Ranking", "Model or optimization service scores distance, workload, capacity, route fit, reliability, and delivery windows", "Improved acceptance, travel time, and on-time delivery"),
    ("Explanation", "Ontology traces the evidence behind each recommendation", "Dispatchers understand and can challenge ranking"),
    ("Shadow evaluation", "Compare recommendations with real assignment and outcome without execution", "Stable uplift across stores and time periods"),
    ("Human-approved execution", "Dispatcher accepts a recommendation through the governed action gateway", "High acceptance and low override/error rate"),
    ("Bounded automation", "Optional auto-assignment only for low-risk cases that satisfy all thresholds", "Audited performance remains within safety limits"),
], [1250, 4850, 3260], font_size=8.5)

heading(doc, "Measurement framework", 1)
add_table(doc, ["Dimension", "Example metrics"], [
    ("Business", "Conversion, basket value, stockout rate, fulfillment cost, on-time delivery, claim resolution time"),
    ("Model", "Precision/recall, ranking gain, ETA error, forecast error, calibration, subgroup performance"),
    ("User", "Recommendation acceptance, override rate, task time, satisfaction, voice correction rate"),
    ("Safety", "Unauthorized tool attempts, policy rejection rate, false automation, PII leakage, grounded-answer rate"),
    ("Platform", "Event lag, graph freshness, reconciliation error, latency, availability, inference cost"),
    ("Learning", "Feedback coverage, experiment velocity, time from model registration to safe deployment"),
], [2000, 7360])

heading(doc, "Risks and mitigations", 1)
add_table(doc, ["Risk", "Consequence", "Mitigation"], [
    ("Ontology mirrors the database", "Little semantic value; tight coupling to legacy schema", "Model business concepts and mappings separately; assign domain owners"),
    ("Dual-write divergence", "Graph and MySQL disagree", "Outbox/CDC, idempotent projections, replay, and reconciliation"),
    ("LLM directly mutates state", "Unauthorized or inconsistent transactions", "Allowlisted governed actions through existing APIs only"),
    ("Stale or over-broad context", "Incorrect answers or privacy leakage", "Freshness metadata, scoped retrieval, tenant filters, and source citations"),
    ("Uncalibrated confidence", "Unsafe automation", "Outcome-based calibration, shadow mode, and impact thresholds"),
    ("Vendor lock-in", "Cost or capability constraints", "Unified model gateway, normalized schemas, portable evaluations"),
    ("Voice transcription error", "Wrong delivery update", "Confidence thresholds, visual transcript, explicit confirmation"),
    ("Excessive platform complexity", "Slow delivery and high operations burden", "Start with a narrow domain graph and managed services where appropriate"),
], [2200, 3000, 4160], font_size=8.6)

heading(doc, "Target technology shape", 1)
add_bullets(doc, [
    "MySQL remains the authoritative operational store.",
    "A transactional outbox and event broker provide reliable semantic synchronization and replay.",
    "A typed property graph stores operational entities, relationships, temporal facts, and provenance.",
    "A vector store supports semantic retrieval of product descriptions, reviews, claims, and operational documents.",
    "A feature store provides point-in-time-correct inputs for ETA, demand, risk, ranking, and anomaly models.",
    "A policy engine and action gateway enforce deterministic constraints before command execution.",
    "A model gateway routes among cloud LLM APIs, private local LLMs, small language models, speech, vision, embeddings, and rerankers.",
    "An evaluation and observability platform connects model inputs and recommendations to approvals, commands, and business outcomes.",
])
add_body(doc, "RDF/OWL is not required for the first implementation. A typed property graph with explicit schema, controlled vocabularies, temporal facts, provenance, and action contracts is likely a better initial fit for operational ecommerce. Standards-based semantic-web technologies can be introduced where interoperability or formal reasoning creates clear value.")

heading(doc, "Decision and next steps", 1)
add_callout(doc, "Target state", "Existing deterministic commerce core + event-driven semantic digital twin + prediction and recommendation services + policy-controlled actions + persona-specific multimodal copilots.", fill=PALE_BLUE, color=BLUE)
add_numbered(doc, [
    "Approve the order-store-inventory-delivery-temperature domain as the first ontology slice.",
    "Name business and technical owners for ontology, events, identity, fulfillment, and AI governance.",
    "Define canonical identifiers, status meanings, temporal semantics, and the first versioned domain events.",
    "Implement a read-only graph projection and reconciliation proof of concept.",
    "Deliver a platform operations query experience and measure answer grounding, freshness, and task-time reduction.",
    "Run ETA and delivery-assignment models in shadow mode before introducing human-approved actions.",
    "Add voice interaction only through the governed action path, beginning with exception reporting and status-update drafts.",
])

heading(doc, "Conclusion", 1)
add_body(doc, "OrderUp already has the transactional foundations and operational data needed for an ontology-based platform. The strategic opportunity is to make the relationships among customers, stores, products, inventory, orders, agents, vehicles, claims, and telemetry explicit and queryable, then place probabilistic intelligence above that semantic foundation.")
add_body(doc, "The architecture succeeds when the ontology improves understanding without becoming a second source of transactional truth, and when AI improves decisions without bypassing business controls. This separation enables innovation in recommendations, forecasting, routing, operations, and voice interaction while preserving the integrity of orders, inventory, payments, identity, and fulfillment.")

# Keep headings from orphaning and set document properties.
doc.core_properties.title = "OrderUp Ontology-Based Platform Case Study"
doc.core_properties.subject = "Future-state AI architecture and migration roadmap"
doc.core_properties.author = "OrderUp Architecture Research"
doc.core_properties.keywords = "OrderUp, ontology, ecommerce, AI architecture, LLM, voice, knowledge graph"

doc.save(OUT)
print(OUT)
